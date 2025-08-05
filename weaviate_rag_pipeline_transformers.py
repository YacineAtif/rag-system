import os
import re
import time
import requests
import json
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from backend.config import Config
from backend.qa_models import ClaudeQA
from backend.llm_generator import LLMGenerator
from neo4j import GraphDatabase
import subprocess
import yaml
import numpy as np
import torch
from sentence_transformers import util
from multi_layer_ood import (
    MultiLayerOODDetector,
    OODDetectionConfig,
    KeywordTiers,
    QualityGates,
    AbstentionConfig,
    ResponseVerificationConfig,
)

os.environ["TOKENIZERS_PARALLELISM"] = "false"

CONFIG = Config()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def load_ood_config() -> OODDetectionConfig:
    try:
        with open("config.yaml", "r") as f:
            data = yaml.safe_load(f).get("ood_detection", {})
    except Exception:
        data = {}

    keywords = KeywordTiers(**data.get("keywords", {}))
    quality = QualityGates(**data.get("quality_gates", {}))
    abstention = AbstentionConfig(**data.get("abstention", {}))
    response = ResponseVerificationConfig(**data.get("response_verification", {}))

    return OODDetectionConfig(
        enabled=data.get("enabled", True),
        similarity_threshold=data.get("similarity_threshold", 0.15),
        graph_connectivity_threshold=data.get("graph_connectivity_threshold", 0.6),
        context_quality_threshold=data.get("context_quality_threshold", 0.7),
        generation_confidence_threshold=data.get("generation_confidence_threshold", 0.8),
        keywords=keywords,
        quality_gates=quality,
        abstention=abstention,
        response_verification=response,
    )

# Haystack v2 imports
from haystack import Document, Pipeline
from haystack.components.writers import DocumentWriter
from haystack.components.embedders import SentenceTransformersTextEmbedder, SentenceTransformersDocumentEmbedder
from haystack_integrations.document_stores.weaviate import WeaviateDocumentStore
from haystack_integrations.components.retrievers.weaviate import WeaviateEmbeddingRetriever

# --- Document Loading from domain_loader.py ---
def load_text_file(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def load_pdf_file(file_path: str) -> str:
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
    except ImportError:
        print("❌ PyPDF2 not installed. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"❌ Error reading PDF {file_path}: {e}")
        return ""

def load_docx_file(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except ImportError:
        print("❌ python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"❌ Error reading DOCX {file_path}: {e}")
        return ""

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        sent_length = len(sentence)
        if current_length + sent_length > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))
            overlap_count = max(1, int(len(current_chunk) * 0.3))
            current_chunk = current_chunk[-overlap_count:]
            current_length = sum(len(s) for s in current_chunk)
        
        current_chunk.append(sentence)
        current_length += sent_length
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


def split_into_sections(text: str, patterns: List[str]) -> List[Tuple[str, str]]:
    """Split text into (section_name, text) using heading patterns."""
    if not patterns:
        return [("unknown", text)]

    lines = text.splitlines()
    sections: List[Tuple[str, List[str]]] = []
    current_name = "unknown"
    buffer: List[str] = []

    compiled = [re.compile(p) for p in patterns]

    for line in lines:
        matched = False
        stripped = line.strip()
        for pat in compiled:
            m = pat.match(stripped)
            if m:
                if buffer:
                    sections.append((current_name, buffer))
                    buffer = []
                current_name = m.group(1).strip().lower()
                matched = True
                break
        if not matched:
            buffer.append(line)

    if buffer:
        sections.append((current_name, buffer))

    return [(name, "\n".join(lines).strip()) for name, lines in sections]

def load_documents_from_folder(folder_path: str) -> List[Document]:
    documents = []
    folder = Path(folder_path)
    
    if not folder.exists():
        print(f"❌ Folder {folder_path} does not exist")
        return documents
    
    print(f"📁 Loading documents from: {folder_path}")
    
    supported_extensions = {
        '.txt': load_text_file,
        '.md': load_text_file,
        '.pdf': load_pdf_file,
        '.docx': load_docx_file,
    }
    
    patterns = CONFIG.chunk_processing.section_patterns

    for file_path in folder.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            print(f"📄 Processing: {file_path.name}")

            try:
                loader_func = supported_extensions[file_path.suffix.lower()]
                content = loader_func(str(file_path))

                if not content.strip():
                    print(f"⚠️  Empty content in {file_path.name}")
                    continue

                file_stats = file_path.stat()
                base_metadata = {
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": file_path.suffix.lower(),
                    "file_size": file_stats.st_size,
                    "modified_date": time.ctime(file_stats.st_mtime)
                }

                sections = split_into_sections(content, patterns)
                chunk_pairs = []
                for section_name, section_text in sections:
                    chs = chunk_text(section_text, chunk_size=500, overlap=100)
                    for ch in chs:
                        chunk_pairs.append((section_name, ch))

                print(f"   Split into {len(chunk_pairs)} chunks")

                for i, (sec_name, chunk) in enumerate(chunk_pairs):
                    chunk_metadata = base_metadata.copy()
                    chunk_metadata.update({
                        "chunk_id": i,
                        "total_chunks": len(chunk_pairs),
                        "content_type": "chunk",
                        "section_name": sec_name
                    })
                    documents.append(Document(content=chunk, meta=chunk_metadata))

            except Exception as e:
                print(f"❌ Error processing {file_path.name}: {e}")
                continue
    
    print(f"✅ Loaded {len(documents)} document chunks")
    return documents

# --- Text Cleaning from simple_qa.py ---
class QueryClassifier:
    """Simple heuristic-based query classifier with config support."""

    def __init__(self, config: Optional[Config] = None):
        self.config = config or CONFIG
        pri_cfg = {}
        sp = getattr(self.config, "section_priorities", {})
        if hasattr(sp, "queries"):
            pri_cfg = sp.queries.get("partnership_queries", {})
        elif isinstance(sp, dict):
            pri_cfg = sp.get("partnership_queries", {})
        self.partnership_keywords = [kw.lower() for kw in pri_cfg.get("keywords", [])]

    def classify(self, query: str) -> str:
        q = query.lower().strip()

        if any(word in q for word in self.partnership_keywords):
            return "partnership"

        entity_keywords = [
            "who",
            "organization",
            "organizations",
            "company",
            "companies",
            "partner",
            "collaborator",
            "team",
            "group",
            "member",
            "participant",
            "contributor",
            "stakeholder",
            "department",
            "division",
            "institution",
        ]

        if any(word in q for word in entity_keywords):
            return "entity"

        if q.startswith("what is") or q.startswith("define") or "definition" in q:
            return "definition"


        procedural_keywords = ["how", "step", "procedure", "process"]
        if any(word in q for word in procedural_keywords):
            return "procedural"

        comparison_keywords = ["compare", "difference", " vs ", " versus "]
        if any(word in q for word in comparison_keywords):
            return "comparison"

        return "general"


class TextProcessor:
    """Utility class for advanced text cleaning and processing."""

    def clean_text(self, text: str, strategy: str = "balanced") -> str:
        if not text:
            return ""

        if strategy == "preserve_structure":
            text = re.sub(r"http[s]?://\S+", "", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()

        if strategy == "clean_moderate":
            text = re.sub(r"http[s]?://\S+", "", text)
            text = re.sub(r"[_*`]+", "", text)
            text = re.sub(r"\s+", " ", text)
            text = re.sub(r"\n+", " ", text)
            return text.strip()

        if strategy == "preserve_lists":
            text = re.sub(r"http[s]?://\S+", "", text)
            text = re.sub(r"[_*`]+", "", text)
            text = re.sub(r"\r\n", "\n", text)
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()

        # balanced strategy
        text = re.sub(r"http[s]?://\S+", "", text)
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"[#*_`<>|]", " ", text)
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n+", " ", text)
        return text.strip()

    def improve_sentence_boundary_detection(self, text: str) -> List[str]:
        text = text.strip()
        if not text:
            return []
        return re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)

    def extract_entities(self, text: str) -> List[str]:
        """Enhanced entity extraction with domain awareness"""
        cleaned = self.clean_text(text, strategy="preserve_structure")
        entities = set()

        # Proper nouns (capitalized phrases)
        entities.update(re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b", cleaned))

        # Acronyms (all caps)
        entities.update(re.findall(r"\b[A-Z]{2,}\b", cleaned))

        # Domain-specific patterns
        entities.update(
            re.findall(r"\b(?:system|project|initiative|platform)\s+[\w-]+\b", cleaned, re.IGNORECASE)
        )

        # Quoted phrases
        entities.update(re.findall(r'"(.*?)"', cleaned))

        common_words = {"the", "and", "for", "with", "this", "that"}
        return [e for e in entities if e and e.lower() not in common_words]

    def preserve_context_formatting(self, text: str) -> str:
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def extract_quality_sentences(self, text: str, strategy: str = "balanced") -> List[str]:
        cleaned_text = self.clean_text(text, strategy=strategy)
        if not cleaned_text:
            return []

        sentences = self.improve_sentence_boundary_detection(cleaned_text)
        quality_sentences = []

        for sentence in sentences:
            sentence = sentence.strip()

            if (
                len(sentence) < 25
                or len(sentence) > 400
                or len(sentence.split()) < 5
                or sentence.lower().startswith(
                    (
                        "what",
                        "how",
                        "when",
                        "where",
                        "why",
                        "is there",
                        "are there",
                    )
                )
                or any(
                    artifact in sentence.lower()
                    for artifact in ["comprehensive guide", "table of contents", "click here"]
                )
                or sentence.count("(") != sentence.count(")")
                or any(artifact in sentence for artifact in ["###", "```", "---", "==="])
            ):
                continue

            sentence = re.sub(r"\s+", " ", sentence).strip()
            if not sentence.endswith((".", "!", "?")):
                sentence += "."

            quality_sentences.append(sentence)

        return quality_sentences


class AnswerGenerator:
    """Generate answers using Claude Haiku."""

    def __init__(self, processor: TextProcessor, config: Optional[Config] = None):
        self.processor = processor
        self.config = config or Config()

    def _score(self, sentence: str, query: str) -> float:
        q_words = set(query.lower().split())
        s_words = set(sentence.lower().split())
        if not q_words:
            return 0.0
        return len(q_words & s_words) / len(q_words)

    def _select_sentences(self, sentences: List[str], query: str, limit: int = 6) -> List[str]:
        scored = [(self._score(s, query), s) for s in sentences]
        scored.sort(key=lambda x: x[0], reverse=True)
        return [s for _, s in scored[:limit] if s]

    def generate(self, query: str, sentences: List[str], query_type: str, history: List[Dict]) -> str:
        if not sentences:
            if history:
                return "I'm not sure. Previously we discussed: " + history[-1]["answer"]
            return "I don't know."

        top_sentences = self._select_sentences(sentences, query)

        instruction = None
        try:
            instruction = self.config.prompting.context_instructions.get(query_type)
        except Exception:
            instruction = None


        claude = ClaudeQA(self.config)
        res = claude.generate(query, top_sentences[:8], instruction=instruction)
        if res.get("answer"):
            return res["answer"]

        if query_type in {"entity", "procedural", "comparison", "general"}:
            if query_type == "entity":
                entities: List[str] = []
                for s in top_sentences:
                    entities.extend(self.processor.extract_entities(s))
                entities = list(dict.fromkeys(entities))
                if entities:
                    return "\n".join([
                        "Entities mentioned:",
                        "- " + "\n- ".join(entities),
                    ])

            if query_type == "procedural":
                steps = [f"{i+1}. {self.processor.clean_text(s, 'preserve_lists')}" for i, s in enumerate(top_sentences)]
                return "Here are the steps:\n" + "\n".join(steps)

            if query_type == "comparison":
                return "Comparison:\n" + "\n".join(f"- {s}" for s in top_sentences[:4])

            if not res.get("answer"):
                res = claude.generate(query, top_sentences[:8], instruction=instruction)
                if res.get("answer"):
                    return res["answer"]


        claude = ClaudeQA(self.config)
        res = claude.generate(query, top_sentences[:8], instruction=instruction)
        return res.get("answer", "")


def create_natural_answer(sentences, query):
    if not sentences:
        return "I don't know."
    
    prefix = "Here's what I found: "
    if len(sentences) == 1:
        return f"{prefix}{sentences[0]}"
    elif len(sentences) == 2:
        return f"{prefix}{sentences[0]} Also, {sentences[1]}"
    else:
        para1 = sentences[0]
        para2 = " ".join(sentences[1:3])  # Limit to avoid too long answers
        return f"{prefix}{para1} Additionally, {para2}"


class Neo4jGraphBuilder:
    def __init__(self, config: Optional[Config] = None):
        self.config = config or CONFIG
        self.driver = GraphDatabase.driver(
            self.config.neo4j.uri,
            auth=(self.config.neo4j.user, self.config.neo4j.password),
        )
        self.claude_extractor = LLMGenerator(model="claude-3-5-haiku-20241022")

    def extract_json_from_claude_response(self, response_text):
        """Extract JSON from Claude's response, handling explanatory text"""
        import re

        response_text = response_text.strip()
        json_patterns = [
            r"```json\s*([\s\S]*?)\s*```",
            r"```\s*([\s\S]*?)\s*```",
            r"\{[\s\S]*\}",
        ]

        for pattern in json_patterns:
            matches = re.findall(pattern, response_text)
            if matches:
                json_text = matches[0] if isinstance(matches[0], str) else matches[0]
                try:
                    return json.loads(json_text)
                except json.JSONDecodeError:
                    continue

        if "Based on" in response_text and "{" in response_text:
            start_idx = response_text.find("{")
            end_idx = response_text.rfind("}") + 1
            if start_idx >= 0 and end_idx > start_idx:
                json_text = response_text[start_idx:end_idx]
                try:
                    return json.loads(json_text)
                except json.JSONDecodeError:
                    pass

        return {"entities": [], "relationships": []}

    def is_valid_chunk(self, text_chunk):
        """Check if chunk has meaningful content for extraction"""
        if not text_chunk or len(text_chunk.strip()) < 20:
            return False
        meaningful_chars = sum(1 for c in text_chunk if c.isalnum())
        if meaningful_chars < len(text_chunk) * 0.3:
            return False
        return True

    def is_graph_populated(self):
        """Check if Neo4j already contains entities"""
        try:
            with self.driver.session() as session:
                result = session.run("MATCH (n) RETURN count(n) as count")
                record = result.single()
                return record["count"] > 0 if record else False
        except Exception:
            return False

    def extract_entities_and_relationships(self, text_chunk, source_doc):
        """Enhanced entity extraction with better error handling"""
        if not self.is_valid_chunk(text_chunk):
            print(f"⏭️  Skipping low-content chunk from {source_doc}")
            return {"entities": [], "relationships": []}

        extraction_prompt = """Extract ALL entities and their relationships from this text comprehensively across any domain:

ENTITIES TO EXTRACT (be thorough and domain-agnostic):
- Organizations (companies, universities, institutions, agencies, groups)
- Projects (research projects, systems, initiatives, programs, studies)
- Technologies (tools, systems, frameworks, methods, platforms)
- People (researchers, engineers, authors, contributors, leaders)
- Concepts (methodologies, theories, approaches, principles)
- Locations (cities, countries, regions, facilities)
- Products (systems, devices, applications, services)

RELATIONSHIPS TO EXTRACT (capture semantic connections):
- Collaboration (PARTNERS_WITH, COLLABORATES_WITH, WORKS_WITH)
- Contribution (CONTRIBUTES_TO, DEVELOPS, CREATES, PROVIDES)
- Participation (PARTICIPATES_IN, LEADS, SUPPORTS, MANAGES)
- Technical (USES, IMPLEMENTS, INTEGRATES_WITH, APPLIES)
- Organizational (EMPLOYS, AFFILIATES_WITH, SPONSORS)
- Location (LOCATED_IN, BASED_IN, OPERATES_IN)

Return comprehensive JSON format:
{{
    "entities": [
        {{"name": "Entity Name", "type": "Organization|Project|Technology|Person|Concept|Location|Product", "properties": {{"description": "brief context"}}}}
    ],
    "relationships": [
        {{"source": "Entity1", "target": "Entity2", "relationship": "RELATIONSHIP_TYPE", "properties": {{"context": "relationship description"}}}}
    ]
}}

Be thorough - extract every meaningful entity and relationship mentioned.

Text: {text}
"""
        try:
            result = self.claude_extractor.generate(
                query=extraction_prompt.format(text=text_chunk),
                context_sentences=[],
                system_prompt="""You are an expert at extracting structured knowledge from text. 

CRITICAL INSTRUCTIONS:
- Return ONLY valid JSON in the exact format specified
- Do NOT include explanatory text before or after the JSON
- If the text contains no meaningful entities, return: {"entities": [], "relationships": []}
- Be comprehensive but ensure valid JSON format

Extract ALL mentioned entities and relationships from the provided text.""",
            )
            extracted = self.extract_json_from_claude_response(result)
            if not isinstance(extracted.get("entities"), list):
                print(f"⚠️  Invalid entities structure in chunk from {source_doc}")
                extracted["entities"] = []
            if not isinstance(extracted.get("relationships"), list):
                print(f"⚠️  Invalid relationships structure in chunk from {source_doc}")
                extracted["relationships"] = []
            entity_count = len(extracted["entities"])
            rel_count = len(extracted["relationships"])
            if entity_count > 0 or rel_count > 0:
                print(f"✅ Extracted {entity_count} entities, {rel_count} relationships from {source_doc}")
            else:
                print(f"ℹ️  No entities/relationships found in chunk from {source_doc}")
            return extracted
        except json.JSONDecodeError as e:
            print(f"❌ JSON parsing failed for chunk from {source_doc}: {e}")
            print(f"Claude response preview: {result[:300]}...")
            return {"entities": [], "relationships": []}
        except Exception as e:
            print(f"❌ Unexpected error in extraction for {source_doc}: {e}")
            return {"entities": [], "relationships": []}

    def populate_graph(self, documents):
        """Phase 1: Build knowledge graph from documents"""
        with self.driver.session() as session:
            for doc in documents:
                extracted = self.extract_entities_and_relationships(doc.content, doc.meta.get("source", ""))

                # Create entities
                for entity in extracted["entities"]:
                    properties = entity.get("properties", {})
                    session.run(
                        "MERGE (e:Entity {name: $name})"
                        " SET e.type = $type, e.source = $source"
                        " SET e += $properties",
                        name=entity["name"],
                        type=entity["type"],
                        source=doc.meta.get("source", ""),
                        properties=properties
                    )

                # Create relationships
                for rel in extracted["relationships"]:
                    rel_type = rel.get("relationship", "RELATES").upper().replace(" ", "_")
                    properties = rel.get("properties", {})
                    query = (
                        f"MATCH (a:Entity {{name: $source}}), (b:Entity {{name: $target}}) "
                        f"MERGE (a)-[r:{rel_type} {{source: $doc_source}}]->(b) "
                        "SET r += $properties"
                    )
                    session.run(
                        query,
                        source=rel["source"],
                        target=rel["target"],
                        doc_source=doc.meta.get("source", ""),
                        properties=properties
                    )

    def validate_graph_population(self):
        """Validate graph population quality (domain-agnostic)"""
        with self.driver.session() as session:
            result = session.run(
                """
                MATCH (n:Entity)
                RETURN n.type as entity_type, count(n) as count
                ORDER BY count DESC
                """
            )
            entity_counts = {record["entity_type"]: record["count"] for record in result}

            rel_result = session.run(
                """
                MATCH ()-[r]->()
                RETURN type(r) as relationship_type, count(r) as count
                ORDER BY count DESC
                """
            )
            relationship_counts = {record["relationship_type"]: record["count"] for record in rel_result}

            total_entities = sum(entity_counts.values())
            total_relationships = sum(relationship_counts.values())

            print(f"📊 Graph validation: {total_entities} entities, {total_relationships} relationships")
            print("📋 Entity types:", entity_counts)
            print("🔗 Relationship types:", relationship_counts)

            if total_entities < 5:
                print("⚠️  Warning: Very few entities extracted. Consider improving extraction prompts.")
            if total_relationships < 3:
                print("⚠️  Warning: Very few relationships extracted. Relationship extraction may need improvement.")
            if len(entity_counts) < 3:
                print("⚠️  Warning: Limited entity type diversity. Extraction may be too narrow.")

    def query_graph(self, query):
        """Enhanced graph querying with intelligent routing"""
        query_lower = query.lower()

        if any(word in query_lower for word in ["contributor", "contribute", "partner", "involved", "collaborate"]):
            return self._query_project_participants(query)
        elif any(word in query_lower for word in ["role", "responsibility", "does", "provide"]):
            return self._query_entity_roles(query)
        elif any(word in query_lower for word in ["develop", "create", "build", "technology"]):
            return self._query_development_relationships(query)
        else:
            return self._query_general_relationships(query)

    def _query_project_participants(self, query):
        """Query for project contributors and partners"""
        with self.driver.session() as session:
            result = session.run(
                """
                MATCH (project:Entity)-[r]-(participant:Entity)
                WHERE project.type = 'Project' OR project.name CONTAINS 'project'
                RETURN participant.name as entity, type(r) as relationship,
                       participant.type as entity_type, project.name as project
                UNION
                MATCH (participant:Entity)-[r:PARTNERS_WITH|CONTRIBUTES_TO|PARTICIPATES_IN]-(other:Entity)
                RETURN participant.name as entity, type(r) as relationship,
                       participant.type as entity_type, other.name as project
                LIMIT 10
                """
            )
            return [
                {"source": record["entity"], "relationship": record["relationship"],
                 "target": record["project"], "type": record["entity_type"]}
                for record in result
            ]

    def _query_entity_roles(self, query):
        """Query for specific entity roles and responsibilities"""
        entity_terms = ["scania", "smart eye", "university", "skövde", "viscando"]
        detected_entity = None
        for term in entity_terms:
            if term in query.lower():
                detected_entity = term
                break

        with self.driver.session() as session:
            if detected_entity:
                result = session.run(
                    """
                    MATCH (entity:Entity)-[r]-(other:Entity)
                    WHERE toLower(entity.name) CONTAINS $entity_term
                    RETURN entity.name as source, type(r) as relationship,
                           other.name as target, entity.properties as properties
                    LIMIT 10
                    """,
                    entity_term=detected_entity
                )
            else:
                result = session.run(
                    """
                    MATCH (entity:Entity)-[r]-(other:Entity)
                    WHERE entity.type = 'Organization'
                    RETURN entity.name as source, type(r) as relationship,
                           other.name as target, entity.properties as properties
                    LIMIT 10
                    """
                )
            return [
                {"source": record["source"], "relationship": record["relationship"],
                 "target": record["target"]}
                for record in result
            ]

    def _query_development_relationships(self, query):
        """Query for development and technology relationships"""
        with self.driver.session() as session:
            result = session.run(
                """
                MATCH (entity:Entity)-[r:DEVELOPS|CREATES|BUILDS|USES|IMPLEMENTS|INTEGRATES_WITH]->(tech:Entity)
                RETURN entity.name as source, type(r) as relationship,
                       tech.name as target, entity.type as entity_type
                LIMIT 10
                """
            )
            return [
                {"source": record["source"], "relationship": record["relationship"],
                 "target": record["target"], "type": record["entity_type"]}
                for record in result
            ]

    def _query_general_relationships(self, query):
        """Fallback general relationship query"""
        with self.driver.session() as session:
            result = session.run(
                """
                MATCH (e1:Entity)-[r]-(e2:Entity)
                WHERE toLower(e1.name) CONTAINS toLower($term) OR toLower(e2.name) CONTAINS toLower($term)
                RETURN e1.name as source, type(r) as relationship, e2.name as target
                LIMIT 10
                """,
                term=query
            )
            return [
                {"source": record["source"], "relationship": record["relationship"], "target": record["target"]}
                for record in result
            ]

    # --- Enhanced Graph Search Strategies ---
    def graph_search(self, query, limit=10):
        """Enhanced graph search with 4-strategy cascade for 100% coverage"""
        try:
            entity_results = self._search_by_entities(query, limit)
        except Exception as e:
            logger.error("Entity search failed: %s", e)
            entity_results = []
        if entity_results:
            logger.info("Graph search using strategy 'entity_match'")
            return entity_results[:limit]

        try:
            relationship_results = self._search_by_relationships(query, limit)
        except Exception as e:
            logger.error("Relationship search failed: %s", e)
            relationship_results = []
        if relationship_results:
            logger.info("Graph search using strategy 'relationship_based'")
            return relationship_results[:limit]

        try:
            content_results = self._search_by_content(query, limit)
        except Exception as e:
            logger.error("Content search failed: %s", e)
            content_results = []
        if content_results:
            logger.info("Graph search using strategy 'content_based'")
            return content_results[:limit]

        try:
            fallback_results = self._get_most_connected_nodes(limit)
        except Exception as e:
            logger.error("Structural fallback search failed: %s", e)
            fallback_results = []
        logger.info("Graph search using strategy 'structural_fallback'")
        return fallback_results[:limit]

    def _search_by_entities(self, query, limit):
        """Primary entity-based search"""
        term = query.lower()
        try:
            with self.driver.session() as session:
                result = session.run(
                    """
                    MATCH (e:Entity)-[r]-(m:Entity)
                    WHERE toLower(e.name) CONTAINS $entity_term
                    RETURN e.name as source, type(r) as relationship, m.name as target
                    LIMIT $limit
                    """,
                    entity_term=term,
                    limit=limit,
                )
                return [
                    {
                        "source": record["source"],
                        "relationship": record["relationship"],
                        "target": record["target"],
                        "relevance_score": 1.0,
                        "search_strategy": "entity_match",
                    }
                    for record in result
                ]
        except Exception as e:
            logger.error("Entity search query failed: %s", e)
            return []

    def _search_by_relationships(self, query, limit):
        """Search graph based on relationship keywords"""
        relationship_mapping = {
            'partner': ['COLLABORATES_WITH', 'PARTNERS_WITH', 'WORKS_WITH'],
            'collaboration': ['COLLABORATES_WITH', 'PARTNERS_WITH'],
            'partnership': ['COLLABORATES_WITH', 'PARTNERS_WITH'],
            'work': ['WORKS_WITH', 'COLLABORATES_WITH'],
            'theory': ['IMPLEMENTS', 'USES', 'APPLIES', 'BASED_ON'],
            'concept': ['IMPLEMENTS', 'DEFINES', 'USES', 'DESCRIBES'],
            'evidence': ['PROVIDES', 'SUPPORTS', 'VALIDATES', 'DEMONSTRATES'],
            'system': ['USES', 'IMPLEMENTS', 'CONTAINS'],
            'project': ['INCLUDES', 'INVOLVES', 'MANAGES']
        }
        query_lower = query.lower()
        relationship_types = set()
        for keyword, rels in relationship_mapping.items():
            if keyword in query_lower:
                relationship_types.update(rels)
        if not relationship_types:
            return []
        try:
            with self.driver.session() as session:
                result = session.run(
                    """
                    MATCH (a)-[r]-(b)
                    WHERE type(r) IN $relationship_types
                    RETURN a.name as source, type(r) as relationship, b.name as target
                    LIMIT $limit
                    """,
                    relationship_types=list(relationship_types),
                    limit=limit,
                )
                return [
                    {
                        "source": record["source"],
                        "relationship": record["relationship"],
                        "target": record["target"],
                        "relevance_score": 0.8,
                        "search_strategy": "relationship_based",
                    }
                    for record in result
                ]
        except Exception as e:
            logger.error("Relationship search query failed: %s", e)
            return []

    def _search_by_content(self, query, limit):
        """Search node properties for query terms"""
        term = query.lower()
        try:
            with self.driver.session() as session:
                result = session.run(
                    """
                    MATCH (n)-[r]-(m)
                    WHERE ANY(prop IN keys(n) WHERE ANY(val IN (CASE WHEN n[prop] IS LIST THEN n[prop] ELSE [n[prop]] END) WHERE toLower(toString(val)) CONTAINS $search_term))
                       OR ANY(prop IN keys(m) WHERE ANY(val IN (CASE WHEN m[prop] IS LIST THEN m[prop] ELSE [m[prop]] END) WHERE toLower(toString(val)) CONTAINS $search_term))
                    RETURN n.name as source, type(r) as relationship, m.name as target
                    LIMIT $limit
                    """,
                    search_term=term,
                    limit=limit,
                )
                return [
                    {
                        "source": record["source"],
                        "relationship": record["relationship"],
                        "target": record["target"],
                        "relevance_score": 0.5,
                        "search_strategy": "content_based",
                    }
                    for record in result
                ]
        except Exception as e:
            logger.error("Content-based search failed: %s", e)
            return []

    def _get_most_connected_nodes(self, limit):
        """Return most connected nodes as structural fallback"""
        try:
            with self.driver.session() as session:
                result = session.run(
                    """
                    MATCH (n)-[r]-(m)
                    WITH n, r, m, COUNT { (n)--() } AS connections
                    RETURN n.name as source, type(r) as relationship, m.name as target, connections
                    ORDER BY connections DESC
                    LIMIT $limit
                    """,
                    limit=limit,
                )
                records = list(result)
                max_conn = max((record["connections"] for record in records), default=1)
                return [
                    {
                        "source": record["source"],
                        "relationship": record["relationship"],
                        "target": record["target"],
                        "relevance_score": record["connections"] / max_conn if max_conn else 0.0,
                        "search_strategy": "structural_fallback",
                    }
                    for record in records
                ]
        except Exception as e:
            logger.error("Structural fallback query failed: %s", e)
            return []

class HybridQueryRouter:

    def __init__(self, graph_builder, vector_retriever, text_embedder):
        self.graph_builder = graph_builder
        self.vector_retriever = vector_retriever
        self.claude = LLMGenerator()
        self.text_embedder = text_embedder
    def should_use_graph(self, query):
        """Determine if query needs graph traversal"""
        graph_keywords = ["partner", "collaborate", "relationship", "connect", "work with", "develop", "contribute"]
        return any(keyword in query.lower() for keyword in graph_keywords)

    def hybrid_retrieve(self, query):
        """Retrieve both vector and graph results for any query."""

        results = {"vector_results": [], "graph_results": []}

        query_embedding = self.text_embedder.run(text=query)["embedding"]
        vector_docs = self.vector_retriever.run(query_embedding=query_embedding)
        results["vector_results"] = vector_docs.get("documents", [])

        try:
            graph_results = self.graph_builder.graph_search(query)
        except Exception as e:
            logger.error("Graph search failed: %s", e)
            graph_results = []
        results["graph_results"] = graph_results

        return results

    def synthesize_answer(self, query, hybrid_results):
        """Combine vector and graph results for final answer"""
        context_parts = []

        # Add vector context
        for doc in hybrid_results["vector_results"][:5]:
            context_parts.append(doc.content)

        # Add graph context
        if hybrid_results["graph_results"]:
            graph_context = "Relationships found: "
            for rel in hybrid_results["graph_results"]:
                graph_context += f"{rel['source']} {rel['relationship']} {rel['target']}. "
            context_parts.append(graph_context)

        # Generate answer using Claude
        return self.claude.generate(
            query=query,
            context_sentences=context_parts,
            system_prompt="Combine information from documents and relationship data to provide comprehensive answers."
        )


def boost_documents(documents: List[Document], query_type: str) -> List[Document]:
    """Apply section-based score boosts to retrieved documents."""
    retrieval_cfg = getattr(CONFIG, "retrieval", {})
    if hasattr(retrieval_cfg, "section_name_boost"):
        base_boost = retrieval_cfg.section_name_boost
    elif isinstance(retrieval_cfg, dict):
        base_boost = retrieval_cfg.get("section_name_boost", 1.0)
    else:
        base_boost = 1.0

    sp = getattr(CONFIG, "section_priorities", {})
    if hasattr(sp, "queries"):
        priority_cfg = sp.queries.get(f"{query_type}_queries", {})
    elif isinstance(sp, dict):
        priority_cfg = sp.get(f"{query_type}_queries", {})
    else:
        priority_cfg = {}
    priority_sections = [s.lower() for s in priority_cfg.get("priority_sections", [])]
    section_factor = priority_cfg.get("boost_factor", 1.0)

    for doc in documents:
        score = doc.score or 0.0
        section = ((doc.meta or {}).get("section_name") or "").lower()
        if section:
            score *= base_boost
            if section in priority_sections:
                score *= section_factor
        doc.score = score

    documents.sort(key=lambda d: d.score or 0.0, reverse=True)
    return documents

class RAGPipeline:
    def __init__(self, document_store, retriever, text_embedder):
        # Existing components
        self.document_store = document_store
        self.retriever = retriever
        self.text_embedder = text_embedder
        # Add Neo4j integration
        self.graph_builder = Neo4jGraphBuilder(CONFIG)
        self.hybrid_router = HybridQueryRouter(self.graph_builder, self.retriever, self.text_embedder)
        self.text_processor = TextProcessor()
        self.ood_detector = MultiLayerOODDetector(load_ood_config())
        self.domain_centroid = self._compute_domain_centroid()

    def _compute_domain_centroid(self):
        """Compute average embedding of all documents as domain centroid."""
        try:
            all_docs = self.document_store.filter_documents()
        except Exception:
            return None
        embeddings = [
            np.array(doc.embedding, dtype=np.float32)
            for doc in all_docs
            if getattr(doc, "embedding", None) is not None
        ]
        if not embeddings:
            return None
        return np.mean(embeddings, axis=0).astype(np.float32)

    def _embedding_similarity(self, query: str) -> float:
        if self.domain_centroid is None:
            return 1.0
        try:
            result = self.text_embedder.run(text=query)
            query_embedding = np.array(result["embedding"], dtype=np.float32)
            return util.pytorch_cos_sim(
                torch.tensor(query_embedding, dtype=torch.float32),
                torch.tensor(self.domain_centroid, dtype=torch.float32),
            ).item()
        except Exception:
            return 1.0

    def get_document_fingerprint(self):
        """Generate fingerprint of all documents to detect changes"""
        import hashlib
        import os

        fingerprint_data = []
        documents_path = Path("documents")

        if not documents_path.exists():
            return ""

        # Get all document files with their modification times and sizes
        for file_path in documents_path.glob("*"):
            if file_path.is_file() and file_path.suffix.lower() in ['.md', '.txt', '.pdf', '.docx']:
                stat = file_path.stat()
                fingerprint_data.append(f"{file_path.name}:{stat.st_mtime}:{stat.st_size}")

        # Create hash of all document metadata
        fingerprint_str = "|".join(sorted(fingerprint_data))
        return hashlib.md5(fingerprint_str.encode()).hexdigest()

    def save_document_fingerprint(self, fingerprint):
        """Save current document fingerprint to file"""
        fingerprint_file = Path(".document_fingerprint")
        with open(fingerprint_file, 'w') as f:
            f.write(fingerprint)

    def load_previous_fingerprint(self):
        """Load previously saved document fingerprint"""
        fingerprint_file = Path(".document_fingerprint")
        if fingerprint_file.exists():
            with open(fingerprint_file, 'r') as f:
                return f.read().strip()
        return ""

    def documents_changed(self):
        """Check if documents have changed since last graph build"""
        current_fingerprint = self.get_document_fingerprint()
        previous_fingerprint = self.load_previous_fingerprint()
        return current_fingerprint != previous_fingerprint

    def is_weaviate_populated(self):
        """Check if Weaviate already contains documents"""
        try:
            existing_docs = self.document_store.filter_documents()
            return len(existing_docs) > 0
        except Exception:
            return False

    def should_process_documents(self):
        """Check if documents need processing"""
        weaviate_populated = self.is_weaviate_populated()
        docs_changed = self.documents_changed()

        if weaviate_populated and not docs_changed:
            return False
        return True

    def process_documents_intelligently(self):
        """Only process documents when necessary"""
        if not self.should_process_documents():
            print("✅ Documents exist in Weaviate and unchanged - skipping document processing")
            print("💡 To force reprocessing, clear Weaviate data or modify documents")
            return

        if self.is_weaviate_populated():
            print("🔄 Documents changed - reprocessing and re-indexing...")
            # Get all documents first, then delete them
            all_docs = self.document_store.filter_documents()
            if all_docs:
                doc_ids = [doc.id for doc in all_docs]
                self.document_store.delete_documents(doc_ids)
            print("🗑️  Cleared existing documents from Weaviate")
        else:
            print("🔄 Processing documents (first run)...")

        print("📁 Loading documents from: documents")
        documents = load_documents_from_folder("documents")
        if not documents:
            print("❌ No documents found to index")
            return

        doc_embedder = SentenceTransformersDocumentEmbedder(
            model="sentence-transformers/all-MiniLM-L6-v2"
        )
        print("🔥 Warming up document embedder...")
        doc_embedder.warm_up()
        print(f"🔄 Embedding {len(documents)} document chunks...")
        embedded_docs = doc_embedder.run(documents)["documents"]
        print(f"🔄 Indexing {len(embedded_docs)} document chunks...")
        indexing_pipeline = Pipeline()
        indexing_pipeline.add_component("writer", DocumentWriter(document_store=self.document_store))
        indexing_pipeline.run({"writer": {"documents": embedded_docs}})
        print("✅ Documents indexed successfully!")

    def populate_knowledge_graph(self):
        """Phase 1: Build knowledge graph (only when needed)"""

        # Check if graph already exists and documents haven't changed
        graph_exists = self.graph_builder.is_graph_populated()
        docs_changed = self.documents_changed()

        if graph_exists and not docs_changed:
            print("✅ Knowledge graph exists and documents unchanged - skipping Phase 1")
            print("💡 To force rebuild, delete Neo4j data or modify documents")
            self.graph_populated = True
            return

        if graph_exists and docs_changed:
            print("🔄 Documents changed - rebuilding knowledge graph...")
            # Optional: Clear existing graph before rebuilding
            print("🗑️  Clearing existing graph data...")
            with self.graph_builder.driver.session() as session:
                session.run("MATCH (n) DETACH DELETE n")
        elif not graph_exists:
            print("🔄 Phase 1: Building Knowledge Graph (first run)...")

        print("🤖 Extracting entities and relationships with Claude...")

        # Get all documents from Weaviate
        all_docs = self.document_store.filter_documents()
        self.graph_builder.populate_graph(all_docs)
        self.graph_builder.validate_graph_population()

        # Save current document fingerprint
        current_fingerprint = self.get_document_fingerprint()
        self.save_document_fingerprint(current_fingerprint)

        print("📊 Knowledge graph populated!")
        self.graph_populated = True

    def force_rebuild_graph(self):
        """Force complete graph rebuild regardless of changes"""
        print("🔄 Force rebuilding knowledge graph...")

        # Clear existing graph
        with self.graph_builder.driver.session() as session:
            session.run("MATCH (n) DETACH DELETE n")

        # Clear fingerprint to force rebuild
        fingerprint_file = Path(".document_fingerprint")
        if fingerprint_file.exists():
            fingerprint_file.unlink()

        # Rebuild graph
        self.populate_knowledge_graph()

    def query_with_graph(self, query):
        """Phase 2: Hybrid query using both vector search and graph traversal"""
        print("🔍 Phase 2: Hybrid retrieval (Vector + Graph)...")

        hybrid_results = self.hybrid_router.hybrid_retrieve(query)

        vector_relevances = [doc.score or 0.0 for doc in hybrid_results["vector_results"]]
        graph_count = len(hybrid_results["graph_results"])
        similarity = self._embedding_similarity(query)
        graph_connectivity = graph_count / max(
            graph_count + len(hybrid_results["vector_results"]), 1
        )

        detection = self.ood_detector.process(
            query=query,
            similarity=similarity,
            graph_connectivity=graph_connectivity,
            retrieved_relevances=vector_relevances,
            token_probs=[0.9],
        )

        if not detection["allow_generation"]:
            return {
                "answer": "I don't have information about this topic.",
                "vector_results": len(hybrid_results["vector_results"]),
                "graph_results": graph_count,
            }

        answer = self.hybrid_router.synthesize_answer(query, hybrid_results)
        sources = [doc.content for doc in hybrid_results["vector_results"]]
        verified, _ = self.ood_detector.verifier.verify(answer, sources, query)
        if not verified:
            answer = "I'm not fully confident about this answer."

        return {
            "answer": answer,
            "vector_results": len(hybrid_results["vector_results"]),
            "graph_results": graph_count,
        }


# --- Main Pipeline ---
def wait_for_weaviate(url="http://localhost:8080", max_retries=30):
    for i in range(max_retries):
        try:
            response = requests.get(f"{url}/v1/.well-known/ready", timeout=5)
            if response.status_code == 200:
                print("✅ Weaviate is ready!")
                return True
        except:
            pass
        print(f"⏳ Waiting for Weaviate... ({i+1}/{max_retries})")
        time.sleep(3)
    return False

def check_docker_containers():
    """Check if required Docker containers are running"""
    try:
        result = subprocess.run(
            ["docker", "ps", "--format", "{{.Names}}"], 
            capture_output=True, text=True, check=True
        )
        running_containers = result.stdout.strip().split('\n')
        
        weaviate_running = any('weaviate' in container for container in running_containers)
        transformers_running = any('transformers' in container for container in running_containers)
        
        return weaviate_running and transformers_running
    except subprocess.CalledProcessError:
        return False

def main():
    print("🤖 Domain-Restricted RAG System (Haystack v2)")
    print("=" * 70)

    import sys
    # Check for force rebuild flag
    force_rebuild = "--rebuild-graph" in sys.argv

    # Start infrastructure if needed
    if check_docker_containers():
        print("✅ Docker containers already running, skipping startup")
    else:
        print("🚀 Starting Docker infrastructure...")
        try:
            subprocess.run(["docker-compose", "up", "-d"], check=True)
            print("✅ Docker containers started")
            time.sleep(30)  # Reduced wait time since containers are starting fresh
        except subprocess.CalledProcessError as e:
            print(f"❌ Failed to start Docker containers: {e}")
            return

    # Wait for Weaviate
    if not wait_for_weaviate():
        print("❌ Weaviate not responding. Check docker-compose logs.")
        return

    # Initialize Weaviate Document Store
    try:
        document_store = WeaviateDocumentStore(url="http://localhost:8080")
        print("✅ Connected to Weaviate")
    except Exception as e:
        print(f"❌ Failed to connect to Weaviate: {e}")
        return

    # Create retriever for hybrid queries
    text_embedder = SentenceTransformersTextEmbedder(
        model="sentence-transformers/all-MiniLM-L6-v2",
    )
    # Warm up the text embedder before any usage
    print("🔥 Warming up text embedder...")
    text_embedder.warm_up()
    retriever = WeaviateEmbeddingRetriever(document_store=document_store)

    # Initialize RAG pipeline with Neo4j integration
    pipeline = RAGPipeline(document_store, retriever, text_embedder)

    # Intelligent document processing
    pipeline.process_documents_intelligently()

    # PHASE 1: Build Knowledge Graph (only when needed)
    if force_rebuild:
        pipeline.force_rebuild_graph()
    else:
        pipeline.populate_knowledge_graph()

    # PHASE 2: Interactive Q&A with hybrid retrieval
    print("\n💬 Hybrid Q&A (Vector + Knowledge Graph)")
    print("Ask questions about partnerships, collaborations, or general content.")

    while True:
        query = input("\n❓ You: ").strip()
        if query.lower() == 'quit':
            break

        start_time = time.time()
        result = pipeline.query_with_graph(query)
        elapsed = time.time() - start_time

        print(f"\n💬 Answer (in {elapsed:.2f}s):")
        print("-" * 60)
        print(result["answer"])
        print(f"\n📊 Retrieved: {result['vector_results']} vector + {result['graph_results']} graph results")
        print("-" * 60)


if __name__ == "__main__":
    main()
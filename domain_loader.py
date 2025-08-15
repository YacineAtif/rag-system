def detect_document_type(text: str, filename: str) -> str:
    """
    Detect document type for content-aware chunking
    Domain-agnostic using configurable indicators from config.yaml
    """
    
    # Get detection config from your config.yaml
    chunking_config = CONFIG.get('chunking', {})
    classification = chunking_config.get('document_classification', {})
    
    # Get configurable indicators
    evidence_indicators = classification.get('evidence_theory_indicators', [])
    evidence_score = sum(1 for indicator in evidence_indicators if indicator.lower() in text.lower())
    
    project_indicators = classification.get('project_indicators', [])
    project_score = sum(1 for indicator in project_indicators if indicator.lower() in text.lower())
    
    safety_indicators = classification.get('safety_indicators', [])
    safety_score = sum(1 for indicator in safety_indicators if indicator.lower() in text.lower())
    
    # Check for high-priority content patterns from config
    content_boost_patterns = CONFIG.get('chunk_processing', {}).get('content_boost_patterns', {})
    high_priority_patterns = content_boost_patterns.get('high_priority', [])
    
    # Apply configurable boosts
    priority_boost = 0
    text_lower = text.lower()
    for pattern in high_priority_patterns:
        try:
            if re.search(pattern, text_lower):
                priority_boost += 2
        except re.error:
            continue
    
    project_score += priority_boost
    
    # Also check filename patterns
    name_lower = filename.lower()
    if any(term in name_lower for term in ['evidence', 'dempster', 'belief']):
        evidence_score += 2
    elif any(term in name_lower for term in ['project', 'application', 'connect']):
        project_score += 2
    elif any(term in name_lower for term in ['traffic', 'safety', 'adas']):
        safety_score += 2
    
    # Determine document type
    if evidence_score > max(project_score, safety_score):
        return 'mathematical'
    elif project_score > max(evidence_score, safety_score):
        return 'project'
    elif safety_score > max(evidence_score, project_score):
        return 'technical'
    else:
        return 'default'#!/usr/bin/env python3
"""
Enhanced Domain Document Loader with Semantic Chunking
Supports: PDF, TXT, DOCX, MD files
Enhanced with content-aware chunking for Evidence Theory, I2Connect, and Safety documents
Special handling for Safety Concept content to ensure proper indexing
"""
import os
import re
import time
import requests
import numpy as np
from pathlib import Path
from typing import List, Dict, Any, Tuple

try:
    import yaml
except Exception:
    yaml = None

CONFIG = {}
if yaml and Path("config.yaml").exists():
    with open("config.yaml", "r") as f:
        CONFIG = yaml.safe_load(f) or {}

from haystack import Document, Pipeline
from haystack.components.writers import DocumentWriter
from haystack.components.embedders import SentenceTransformersDocumentEmbedder
from haystack_integrations.document_stores.weaviate import WeaviateDocumentStore
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer

def load_text_file(file_path: str) -> str:
    """Load content from a text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def load_pdf_file(file_path: str) -> str:
    """Load content from a PDF file"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
    except ImportError:
        print("⚠ PyPDF2 not installed. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"⚠ Error reading PDF {file_path}: {e}")
        return ""

def load_docx_file(file_path: str) -> str:
    """Load content from a DOCX file"""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except ImportError:
        print("⚠ python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"⚠ Error reading DOCX {file_path}: {e}")
        return ""

def detect_document_type(text: str, filename: str) -> str:
    """
    Detect document type for content-aware chunking
    Domain-agnostic using configurable indicators from config.yaml
    """
    
    # Get detection config from your config.yaml
    chunking_config = CONFIG.get('chunking', {})
    classification = chunking_config.get('document_classification', {})
    
    # Get configurable indicators
    evidence_indicators = classification.get('evidence_theory_indicators', [])
    evidence_score = sum(1 for indicator in evidence_indicators if indicator.lower() in text.lower())
    
    project_indicators = classification.get('project_indicators', [])
    project_score = sum(1 for indicator in project_indicators if indicator.lower() in text.lower())
    
    safety_indicators = classification.get('safety_indicators', [])
    safety_score = sum(1 for indicator in safety_indicators if indicator.lower() in text.lower())
    
    # Check for high-priority content patterns from config
    content_boost_patterns = CONFIG.get('chunk_processing', {}).get('content_boost_patterns', {})
    high_priority_patterns = content_boost_patterns.get('high_priority', [])
    
    # Apply configurable boosts
    priority_boost = 0
    text_lower = text.lower()
    for pattern in high_priority_patterns:
        try:
            if re.search(pattern, text_lower):
                priority_boost += 2
        except re.error:
            continue
    
    project_score += priority_boost
    
    # Also check filename patterns
    name_lower = filename.lower()
    if any(term in name_lower for term in ['evidence', 'dempster', 'belief']):
        evidence_score += 2
    elif any(term in name_lower for term in ['project', 'application', 'connect']):
        project_score += 2
    elif any(term in name_lower for term in ['traffic', 'safety', 'adas']):
        safety_score += 2
    
    # Determine document type
    if evidence_score > max(project_score, safety_score):
        return 'mathematical'
    elif project_score > max(evidence_score, safety_score):
        return 'project'
    elif safety_score > max(evidence_score, project_score):
        return 'technical'
    else:
        return 'default'

def get_chunking_strategy(doc_type: str) -> Dict[str, Any]:
    """
    Get chunking configuration for document type
    Uses your config.yaml enhanced chunking settings
    """
    
    chunking_config = CONFIG.get('chunking', {})
    content_specific = chunking_config.get('content_specific', {})
    
    if doc_type in content_specific:
        strategy = content_specific[doc_type].copy()
        # Add defaults from main config
        strategy.setdefault('chunk_size', chunking_config.get('chunk_size', CONFIG.get('chunk_size', 800)))
        strategy.setdefault('chunk_overlap', chunking_config.get('chunk_overlap', CONFIG.get('chunk_overlap', 150)))
        return strategy
    
    # Fallback to your existing chunking config
    return {
        'chunk_size': chunking_config.get('chunk_size', CONFIG.get('chunk_size', 800)),
        'chunk_overlap': chunking_config.get('chunk_overlap', CONFIG.get('chunk_overlap', 150))
    }

def find_safety_concept_boundaries(text: str) -> List[int]:
    """
    Find Safety Concept section boundaries to preserve them intact
    """
    boundaries = []
    
    # Patterns specifically for Safety Concept sections
    safety_patterns = [
        r'###?\s*\d+\.\d+\s+.*?(?:Safety\s+)?Concept.*?,
        r'###?\s*.*?Concept\s+\d+.*?,
        r'###?\s*Safety\s+Concept.*?,
        r'###?\s*\d+\.\d+\s+Comparative.*?,
        r'\n\s*\*\*Safety\s+Concept.*?\*\*',
        r'\|\s*Feature.*?\|\s*Concept\s+1.*?\|',  # Table headers
    ]
    
    for pattern in safety_patterns:
        try:
            for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
                boundaries.append(match.start())
        except re.error:
            continue
    
    return boundaries

def find_section_boundaries(text: str, doc_type: str) -> List[int]:
    """
    Find good places to break chunks based on document structure
    Domain-agnostic with configurable high-priority content detection
    """
    
    boundaries = []
    chunking_config = CONFIG.get('chunking', {})
    section_patterns = chunking_config.get('section_patterns', {})
    
    # Get patterns for this document type
    high_priority = section_patterns.get('high_priority', [])
    medium_priority = section_patterns.get('medium_priority', [])
    
    # PRIORITY: Find high-priority content boundaries first
    priority_boundaries = find_high_priority_content_boundaries(text)
    boundaries.extend(priority_boundaries)
    
    # Find high priority boundaries (must preserve these sections)
    for pattern in high_priority:
        try:
            for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
                boundaries.append(match.start())
        except re.error:
            continue  # Skip invalid patterns
    
    # Find medium priority boundaries (good places to break)
    for pattern in medium_priority:
        try:
            for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
                boundaries.append(match.start())
        except re.error:
            continue
    
    # Add paragraph boundaries as low priority
    for match in re.finditer(r'\n\s*\n', text):
        boundaries.append(match.start())
    
    # Sort and deduplicate boundaries
    boundaries = sorted(list(set(boundaries)))
    return boundaries

def find_best_boundary(boundaries: List[int], target: int, min_pos: int, max_pos: int) -> int:
    """Find the best boundary near the target position"""
    
    # Look for boundaries within reasonable distance of target
    tolerance = target * 0.2  # 20% tolerance
    
    candidates = [b for b in boundaries if min_pos <= b <= max_pos and abs(b - target) <= tolerance]
    
    if candidates:
        # Return the boundary closest to target
        return min(candidates, key=lambda x: abs(x - target))
    
    # No good boundary found
    return target

def enhanced_chunk_text(text: str, doc_type: str = 'default', filename: str = '') -> List[str]:
    """
    Enhanced version of chunk_text function with content-aware chunking
    Domain-agnostic with configurable high-priority content handling
    """
    
    # Get chunking strategy for this document type
    chunk_config = get_chunking_strategy(doc_type)
    chunk_size = chunk_config['chunk_size']
    overlap = chunk_config['chunk_overlap']
    
    # Check for high-priority content patterns from config
    content_boost_patterns = CONFIG.get('chunk_processing', {}).get('content_boost_patterns', {})
    high_priority_patterns = content_boost_patterns.get('high_priority', [])
    
    # Detect high-priority content and adjust chunking accordingly
    has_high_priority_content = False
    text_lower = text.lower()
    for pattern in high_priority_patterns:
        try:
            if re.search(pattern, text_lower):
                has_high_priority_content = True
                break
        except re.error:
            continue
    
    # Optimize chunking for high-priority content
    if has_high_priority_content:
        chunk_size = min(chunk_size, 700)  # Smaller chunks for better precision
        overlap = max(overlap, 100)  # More overlap for better context
        print(f"   🎯 High-priority content detected - using optimized chunking (size: {chunk_size}, overlap: {overlap})")
    
    # Your existing bullet list logic (preserved)
    bullet_re = re.compile(r"^\s*(?:[-*]|\d+\.)\s+")
    lines = text.splitlines()
    segments: List[str] = []
    buffer: List[str] = []
    in_bullet = False

    for line in lines:
        if bullet_re.match(line):
            if not in_bullet and buffer:
                segments.append(" ".join(buffer).strip())
                buffer = []
            in_bullet = True
            buffer.append(line.strip())
        else:
            if in_bullet and buffer:
                segments.append("\n".join(buffer).strip())
                buffer = []
            in_bullet = False
            buffer.append(line.strip())

    if buffer:
        if in_bullet:
            segments.append("\n".join(buffer).strip())
        else:
            segments.append(" ".join(buffer).strip())

    # Enhanced chunking with section boundary awareness
    chunks: List[str] = []
    current_chunk: List[str] = []
    current_tokens = 0
    
    # Find section boundaries for smarter chunking
    boundaries = find_section_boundaries(text, doc_type)

    for seg in segments:
        seg_tokens = len(seg.split())
        if current_tokens + seg_tokens > chunk_size and current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append(chunk_text)
            
            if overlap > 0:
                tokens = chunk_text.split()
                overlap_tokens = tokens[-overlap:]
                current_chunk = [" ".join(overlap_tokens)]
                current_tokens = len(overlap_tokens)
            else:
                current_chunk = []
                current_tokens = 0
        current_chunk.append(seg)
        current_tokens += seg_tokens

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks

def calculate_chunk_metadata(chunk_text: str, doc_type: str, filename: str) -> Dict[str, Any]:
    """
    Calculate enhanced metadata for chunks
    Domain-agnostic with configurable content pattern detection
    """
    
    metadata = {}
    
    # Document type and technical analysis
    metadata['document_type'] = doc_type
    metadata['chunk_size_words'] = len(chunk_text.split())
    
    # Get configurable content markers from config
    content_markers = CONFIG.get('chunk_processing', {}).get('content_markers', {})
    comparative_markers = content_markers.get('comparative_content', ['comparative', 'comparison', 'vs', 'versus'])
    technical_markers = content_markers.get('technical_content', ['system', 'module', 'component'])
    conceptual_markers = content_markers.get('conceptual_content', ['concept', 'theory', 'framework'])
    
    # Content analysis using configurable patterns
    chunk_lower = chunk_text.lower()
    metadata['contains_equations'] = bool(re.search(r'(?i)(bel\(|pl\(|m\(|∑|∫|∀|∃|theorem|proof)', chunk_text))
    metadata['contains_work_packages'] = bool(re.search(r'(?i)(work package|wp\s*\d+|task \d+\.\d+|deliverable)', chunk_text))
    metadata['contains_technical_terms'] = any(marker in chunk_lower for marker in technical_markers)
    metadata['contains_comparative_content'] = any(marker in chunk_lower for marker in comparative_markers)
    metadata['contains_conceptual_content'] = any(marker in chunk_lower for marker in conceptual_markers)
    
    # Check for high-priority patterns from config
    content_boost_patterns = CONFIG.get('chunk_processing', {}).get('content_boost_patterns', {})
    high_priority_patterns = content_boost_patterns.get('high_priority', [])
    
    # Calculate content relevance score
    relevance_score = 0
    for pattern in high_priority_patterns:
        try:
            if re.search(pattern, chunk_lower):
                relevance_score += 1
        except re.error:
            continue
    
    metadata['high_priority_relevance'] = min(relevance_score / max(len(high_priority_patterns), 1), 1.0)
    
    # Technical density calculation (configurable terms from entities in config)
    metadata_config = CONFIG.get('metadata_extraction', {})
    entities_config = metadata_config.get('entities', {})
    
    technical_terms = []
    for entity_type, terms in entities_config.items():
        technical_terms.extend(terms)
    
    # Fallback technical terms if config is empty
    if not technical_terms:
        technical_terms = ['ADAS', 'Evidence Theory', 'Risk Assessment', 'Driver Monitoring']
    
    words = chunk_text.split()
    if words:
        technical_count = sum(1 for word in words if any(term.lower() in word.lower() for term in technical_terms))
        metadata['technical_density'] = min(technical_count / len(words) * 10, 1.0)
    else:
        metadata['technical_density'] = 0.0
    
    # Extract key entities from config
    metadata['key_entities'] = []
    
    for entity_type, entities in entities_config.items():
        for entity in entities:
            if entity.lower() in chunk_text.lower():
                metadata['key_entities'].append(entity)
    
    return metadata

def split_into_sections(text: str, patterns: List[str]) -> List[Tuple[str, str]]:
    """
    Split text into (section_name, text) using heading patterns
    Enhanced to preserve Safety Concept sections
    """
    if not patterns:
        return [("unknown", text)]

    lines = text.splitlines()
    sections: List[Tuple[str, List[str]]] = []
    current_name = "unknown"
    buffer: List[str] = []

    compiled = [re.compile(p) for p in patterns]

    for line in lines:
        matched = False
        stripped = line.strip()
        for pat in compiled:
            try:
                m = pat.match(stripped)
                if m:
                    if buffer:
                        sections.append((current_name, buffer))
                        buffer = []
                    
                    # Extract section name and detect Safety Concepts
                    section_name = m.group(1).strip().lower()
                    
                    # Special handling for Safety Concept sections
                    if 'safety concept' in section_name or 'concept 1' in section_name or 'concept 2' in section_name:
                        # Preserve the original case for Safety Concept sections
                        section_name = m.group(1).strip()
                    
                    current_name = section_name
                    matched = True
                    break
            except (re.error, IndexError):
                continue  # Skip invalid patterns
        if not matched:
            buffer.append(line)

    if buffer:
        sections.append((current_name, buffer))

    return [(name, "\n".join(lines).strip()) for name, lines in sections]

def load_documents_from_folder(folder_path: str) -> List[Document]:
    """
    Load all supported documents from a folder
    Enhanced with content-aware chunking and Safety Concept support
    """
    documents = []
    folder = Path(folder_path)
    
    if not folder.exists():
        print(f"⚠ Folder {folder_path} does not exist")
        return documents
    
    print(f"📁 Loading documents from: {folder_path}")
    
    supported_extensions = {
        '.txt': load_text_file,
        '.md': load_text_file,
        '.pdf': load_pdf_file,
        '.docx': load_docx_file,
    }
    
    # Get section patterns from config
    patterns = CONFIG.get("chunk_processing", {}).get("section_patterns", [])

    for file_path in folder.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            print(f"📄 Processing: {file_path.name}")

            try:
                loader_func = supported_extensions[file_path.suffix.lower()]
                content = loader_func(str(file_path))

                if not content.strip():
                    print(f"⚠️ Empty content in {file_path.name}")
                    continue

                # Detect document type for enhanced chunking
                doc_type = detect_document_type(content, file_path.name)
                print(f"   Document type: {doc_type}")
                
                # Check for high-priority content using configurable patterns
                content_boost_patterns = CONFIG.get('chunk_processing', {}).get('content_boost_patterns', {})
                high_priority_patterns = content_boost_patterns.get('high_priority', [])
                
                has_high_priority_content = False
                content_lower = content.lower()
                for pattern in high_priority_patterns:
                    try:
                        if re.search(pattern, content_lower):
                            has_high_priority_content = True
                            break
                    except re.error:
                        continue
                
                if has_high_priority_content:
                    print(f"   🎯 High-priority content detected!")

                file_stats = file_path.stat()
                base_metadata = {
                    "filename": file_path.name,
                    "file_path": str(file_path),
                    "file_type": file_path.suffix.lower(),
                    "file_size": file_stats.st_size,
                    "modified_date": time.ctime(file_stats.st_mtime),
                    "document_type": doc_type,
                    "has_high_priority_content": has_high_priority_content
                }

                sections = split_into_sections(content, patterns)
                chunk_pairs = []
                
                for section_name, section_text in sections:
                    # Use enhanced chunking instead of basic chunk_text
                    chunks = enhanced_chunk_text(section_text, doc_type, file_path.name)
                    
                    for chunk in chunks:
                        chunk_pairs.append((section_name, chunk))

                print(f"   Split into {len(chunk_pairs)} chunks using {doc_type} strategy")
                
                # Log high-priority content chunks for verification
                high_priority_chunks = 0
                for _, chunk in chunk_pairs:
                    # Use the same patterns to check chunks
                    chunk_lower = chunk.lower()
                    for pattern in high_priority_patterns:
                        try:
                            if re.search(pattern, chunk_lower):
                                high_priority_chunks += 1
                                break
                        except re.error:
                            continue
                
                if high_priority_chunks > 0:
                    print(f"   🎯 {high_priority_chunks} chunks contain high-priority content")

                for i, (sec_name, chunk) in enumerate(chunk_pairs):
                    chunk_metadata = base_metadata.copy()
                    chunk_metadata.update({
                        "chunk_id": i,
                        "total_chunks": len(chunk_pairs),
                        "content_type": "chunk",
                        "section_name": sec_name
                    })
                    
                    # Add enhanced metadata with Safety Concept detection
                    enhanced_metadata = calculate_chunk_metadata(chunk, doc_type, file_path.name)
                    chunk_metadata.update(enhanced_metadata)
                    
                    # Log high-priority content chunks for debugging
                    if enhanced_metadata.get('high_priority_relevance', 0) > 0:
                        print(f"      🎯 Chunk {i}: High-priority relevance = {enhanced_metadata['high_priority_relevance']:.2f}")

                    documents.append(Document(content=chunk, meta=chunk_metadata))

            except Exception as e:
                print(f"⚠ Error processing {file_path.name}: {e}")
                continue
    
    # Final verification
    high_priority_docs = [doc for doc in documents if doc.meta.get('high_priority_relevance', 0) > 0]
    print(f"✅ Loaded {len(documents)} document chunks")
    print(f"🎯 {len(high_priority_docs)} chunks contain high-priority content")
    
    return documents

def wait_for_weaviate(url="http://localhost:8080", max_retries=10):
    """Wait for Weaviate to be ready"""
    for i in range(max_retries):
        try:
            response = requests.get(f"{url}/v1/meta")
            if response.status_code == 200:
                print("✅ Weaviate is ready!")
                return True
        except:
            pass
        print(f"⏳ Waiting for Weaviate... ({i+1}/{max_retries})")
        time.sleep(2)
    return False

def setup_domain_knowledge_base(documents_folder: str = "documents"):
    """Set up the complete domain knowledge base with enhanced chunking"""
    print("🚀 Domain Knowledge Loader with Enhanced Semantic Chunking")
    print("🎯 Configurable High-Priority Content Support")
    print("=" * 70)
    
    if not wait_for_weaviate():
        print("⚠ Weaviate not responding. Make sure docker-compose is running.")
        return None, None
    
    documents = load_documents_from_folder(documents_folder)
    
    if not documents:
        print("⚠ No documents found to index")
        return None, None
    
    try:
        print("🔧 Initializing Weaviate document store...")
        document_store = WeaviateDocumentStore(url="http://localhost:8080")
        
        print("📚 Setting up indexing pipeline...")
        indexing_pipeline = Pipeline()
        indexing_pipeline.add_component(
            "embedder",
            SentenceTransformersDocumentEmbedder(
                model="sentence-transformers/all-MiniLM-L6-v2",
                progress_bar=False,
                encode_kwargs={"show_progress_bar": False, "convert_to_tensor": False},
            )
        )
        indexing_pipeline.add_component("writer", DocumentWriter(document_store=document_store))
        indexing_pipeline.connect("embedder", "writer")
        
        print(f"🔥 Indexing {len(documents)} document chunks...")
        print("   Enhanced features:")
        print("   - Content-aware chunking (Evidence Theory, Project, Technical)")
        print("   - Section boundary preservation")
        print("   - Enhanced metadata extraction")
        print("   🎯 Configurable high-priority content optimization")
        
        indexing_pipeline.run({"embedder": {"documents": documents}})
        print("✅ Domain documents indexed successfully with enhanced chunking!")
        
        # Verify high-priority content indexing
        high_priority_docs = [doc for doc in documents if doc.meta.get('high_priority_relevance', 0) > 0]
        print(f"🎯 Verification: {len(high_priority_docs)} high-priority content chunks indexed")
        
        return document_store
        
    except Exception as e:
        print(f"⚠ Error setting up knowledge base: {e}")
        return None

if __name__ == "__main__":
    if not os.path.exists("documents"):
        print("⚠ 'documents' folder not found. Please create it and add your domain documents.")
    else:
        setup_domain_knowledge_base("documents")